from contextlib import contextmanager
import win32com.client as win32, pywintypes
from general.fileutil import path_with_suffix
from general.singleton import Singleton
from docx import Document

class WordReaderException(Exception):pass

class DocxWordDocument:
    def __init__(self, doc_path=None):
        self.doc_path = doc_path
        self._document = Document(doc_path)
    @property 
    def document(self):
        return self._document
    @contextmanager
    def open_document(self, doc_path):
        if self.document:
            self._close()
        self._document = Document(doc_path)
        self.doc_path = doc_path
        print('a')
        yield self
        self._close()
    def save(self):        
        if self.document:
            self.document.save(self.doc_path)
        return self.doc_path
    def find_table(self, index=1):
        if self.document and len(self.document.tables) >= index:
            return self.document.tables[index-1]
        else:
            return None
    def _close(self):
        self._document = None
        self.doc_path = ''



class WordApplication(Singleton):
    def __init__(self):
        self.word= win32.dynamic.Dispatch('word.application')
        self.word.visible = 0

wdFormatFilteredHTML = 10
wdFormatPDF = 17
wdDoNotSaveChanges = 0
class WordDocument:
    def __init__(self, doc_path=None):
        self._word = WordApplication()
        self.doc_path = doc_path
        self._document = None
    @property
    def word(self):
        return self._word.word
    @property 
    def document(self):
        return self._document
    @contextmanager
    def open_document(self, doc_path, read_only = True):
        if self.document:
            self.close()
        if read_only:
            self.word.Documents.Open(doc_path, ReadOnly=-1)
        else:
            self.word.Documents.Open(doc_path)
        self._document = self.word.ActiveDocument
        self.doc_path = doc_path
        yield self
        self._close()
    def _save_as(self, file_format, suffix, filename=None):        
        save_name = str(filename) if filename else str(path_with_suffix(self.doc_path, suffix))
        self.document.SaveAs(save_name, FileFormat=file_format)
        return save_name
    def save_as_pdf(self, filename=None):
        return self._save_as(wdFormatPDF, '.pdf', filename=filename)
    def save_as_htm(self, filename=None):
        return self._save_as(wdFormatFilteredHTML, '.htm', filename=filename)
    def find_table(self, index=1):
        if self.document and self.document.Tables.Count >= index:
            return self.document.Tables(index)
        else:
            return None
    def _close(self):
        try:
            if self.document:
                self.document.Close(SaveChanges=wdDoNotSaveChanges)
            self._document = None
            self.doc_path = ''
        except pywintypes.com_error as E:
            pass 
        # the COM system sometimes seems to close itself when not necessary, 
        # 'The object invoked has disconnected from its clients.'
        # it seems harmless to ignore this.
        #TODO: research and solve

            