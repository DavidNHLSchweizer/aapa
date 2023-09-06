from contextlib import contextmanager
from general.fileutil import path_with_suffix
from docx import Document
import docx2pdf 
from pythoncom import CoInitialize

class WordReaderException(Exception):pass

class Word2PdfConvertor:
    def convert(self, input_file, output_file, keep_active=False):
        docx2pdf.convert(input_file, output_file, keep_active)  
        return output_file               

class DocxWordDocument:
    def __init__(self, doc_path=None):
        self.doc_path = doc_path
        CoInitialize() #needed for TUI interface, else async or multithreading problems
        self._document = Document(doc_path)
        self.pdf_convertor = Word2PdfConvertor()
        self._modified = False
    @property 
    def document(self):
        return self._document
    @property
    def modified(self):
        return self._modified
    @contextmanager
    def open_document(self, doc_path):
        if self.document:
            self._close()
        self._document = Document(doc_path)
        self.doc_path = doc_path
        yield self
        self._close()
    def save(self):        
        if self.document:
            self.document.save(self.doc_path)
        return self.doc_path
    def save_as_pdf(self, pdf_name, last_doc=False):
        if self.doc_path:
            return self.pdf_convertor.convert(self.doc_path, path_with_suffix(pdf_name, '.pdf'), not last_doc)
    def find_table(self, index=1):
        if self.document and len(self.document.tables) >= index:
            return self.document.tables[index-1]
        else:
            return None
    def read_table_cell(self, table, cell_row, cell_col)->str:
        if not table:
            return ''
        return table.cell(cell_row-1,cell_col-1).text
    def write_table_cell(self, table, cell_row, cell_col, value):
        if not table:
            return
        old_value = table.cell(cell_row-1, cell_col-1).text
        table.cell(cell_row-1, cell_col-1).text = value
        self._modified = old_value != value
    def _close(self):
        self._document = None
        self.doc_path = ''

